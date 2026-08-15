#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成两份用于「导入公众号 / 头条」验证的样例 Word 文档（.docx）。

每份文档都覆盖 Markdown 的常见格式：
  标题(H1/H2/H3) / 加粗 / 斜体 / 行内代码 / 有序列表 / 无序列表
  / 任务清单(勾选框) / 引用 / 表格 / 代码块 / 图片 / 公式(文本形式)

图片用 Pillow 现场生成（带文字的彩色图），直接内嵌进 docx，
导入平台时无需额外准备图片文件。

输出目录：本脚本同级的 ../sample-articles/
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from PIL import Image, ImageDraw, ImageFont

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(HERE, "..", "sample-articles")
os.makedirs(OUT_DIR, exist_ok=True)


# --------------------------------------------------------------------------
# 生成示例图片（带文字的彩色图，模拟封面 / 配图）
# --------------------------------------------------------------------------
def _font(size):
    candidates = [
        "C:/Windows/Fonts/msyh.ttc",      # 微软雅黑
        "C:/Windows/Fonts/simhei.ttf",    # 黑体
        "C:/Windows/Fonts/simsun.ttc",    # 宋体
    ]
    for c in candidates:
        if os.path.exists(c):
            try:
                return ImageFont.truetype(c, size)
            except Exception:
                pass
    return ImageFont.load_default()


def make_image(path, w, h, bg, title, subtitle):
    img = Image.new("RGB", (w, h), bg)
    d = ImageDraw.Draw(img)
    f_title = _font(46)
    f_sub = _font(24)
    # 标题
    bbox = d.textbbox((0, 0), title, font=f_title)
    tw = bbox[2] - bbox[0]
    d.text(((w - tw) // 2, int(h * 0.34)), title, font=f_title, fill=(255, 255, 255))
    # 副标题
    bbox2 = d.textbbox((0, 0), subtitle, font=f_sub)
    sw = bbox2[2] - bbox2[0]
    d.text(((w - sw) // 2, int(h * 0.52)), subtitle, font=f_sub, fill=(235, 235, 235))
    # 底部水印条
    d.rectangle([0, h - 40, w, h], fill=(0, 0, 0))
    d.text((16, h - 30), "MarkItDown Desktop · 示例配图", font=_font(18), fill=(200, 200, 200))
    img.save(path)


def gen_images(prefix):
    cover = os.path.join(OUT_DIR, f"{prefix}-cover.png")
    diagram = os.path.join(OUT_DIR, f"{prefix}-diagram.png")
    make_image(cover, 960, 420, (37, 99, 175), "示例封面图", "Cover Image (embedded)")
    make_image(diagram, 960, 420, (22, 138, 96), "示例配图", "Inline Image (embedded)")
    return cover, diagram


# --------------------------------------------------------------------------
# docx 小工具
# --------------------------------------------------------------------------
def set_code_shading(paragraph):
    """给段落加浅灰底纹，模拟代码块背景。"""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    set_code_shading(p)
    for i, line in enumerate(code.split("\n")):
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10.5)
        # 中文字体回退
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        if i != len(code.split("\n")) - 1:
            run.add_break()
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def add_checklist(doc, items):
    for text, done in items:
        mark = "☑" if done else "☐"
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{mark} {text}")
        if done:
            r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            r.font.strike = True
    return


# --------------------------------------------------------------------------
# 文章一：公众号（Markdown 写作）
# --------------------------------------------------------------------------
def build_wechat():
    cover, diagram = gen_images("wechat")
    doc = Document()
    doc.add_heading("用 Markdown 高效写作：从入门到发布公众号", level=1)
    doc.add_paragraph(
        "本文 demo 覆盖公众号导入文档支持的主要格式，用于验证「导出 .docx → 公众号导入」"
        "的内容还原度。加粗、**重点**、行内代码 `npm run build` 都能保留。"
    )

    doc.add_heading("一、为什么用 Markdown", level=2)
    doc.add_paragraph("Markdown 让你专注内容，不被排版打断：")
    doc.add_paragraph("写作时只管结构，发布前再统一美化样式。", style="List Bullet")
    doc.add_paragraph("纯文本存储，跨平台、跨工具都能打开。", style="List Bullet")
    doc.add_paragraph("可与图床、CI、静态站点无缝衔接。", style="List Bullet")

    doc.add_heading("二、常用语法速查", level=2)
    doc.add_paragraph("下面是有序列表示例（步骤型内容）：")
    doc.add_paragraph("安装 Node 与编辑器", style="List Number")
    doc.add_paragraph("新建 .md 文件并写入标题", style="List Number")
    doc.add_paragraph("导出为 .docx 后导入公众号", style="List Number")

    doc.add_heading("三、任务清单（勾选框）", level=3)
    add_checklist(doc, [
        ("确定文章主题与受众", True),
        ("列出大纲与小标题", True),
        ("填充正文并配图", False),
        ("预览手机端排版效果", False),
    ])

    doc.add_heading("四、引用与表格", level=2)
    add_quote(doc, "好的排版，是让内容自己说话。")
    doc.add_paragraph("常见格式支持对照：")
    table = doc.add_table(rows=5, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "格式", "公众号导入", "备注"
    rows = [
        ("标题 / 加粗 / 斜体", "支持", "还原度好"),
        ("表格", "支持", "可能微调列宽"),
        ("代码块", "基本支持", "可能转为图片"),
        ("公式 / 流程图", "有限", "可能不显示，见下"),
    ]
    for i, (a, b, c) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text, cells[1].text, cells[2].text = a, b, c

    doc.add_heading("五、代码块示例", level=2)
    add_code_block(doc,
                   "# 导出为公众号 Word 文档\n"
                   "npm run build\n"
                   "选择「导出 → 公众号 (.docx)」\n"
                   "在公众号编辑页「导入文档」选该文件")

    doc.add_heading("六、配图（图片）", level=2)
    doc.add_paragraph("封面图：")
    doc.add_picture(cover, width=Inches(4.6))
    doc.add_paragraph("正文配图：")
    doc.add_picture(diagram, width=Inches(4.6))

    doc.add_heading("七、关于公式", level=2)
    doc.add_paragraph(
        "公众号导入对数学公式（KaTeX / LaTeX）支持有限，可能转为图片或直接不显示。"
        "下方以纯文本给出示例，导入后请确认是否保留：质能方程 E = mc²；"
        "勾股定理 a² + b² = c²。"
    )

    doc.add_paragraph("—— 完 ——")
    out = os.path.join(OUT_DIR, "公众号样例-Markdown高效写作.docx")
    doc.save(out)
    return out



if __name__ == "__main__":
    a = build_wechat()
    print("生成完成：")
    print(" ", os.path.abspath(a))
